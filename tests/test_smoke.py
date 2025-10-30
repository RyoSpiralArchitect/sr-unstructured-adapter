from __future__ import annotations

import base64
import json
import zipfile
from pathlib import Path

import pytest

_docx_pkg = pytest.importorskip(
    "docx", reason="python-docx is required for DOCX smoke tests"
)
_openpyxl_pkg = pytest.importorskip(
    "openpyxl", reason="openpyxl is required for XLSX smoke tests"
)

from docx import Document as DocxDocument
from openpyxl import Workbook

from sr_adapter import build_payload, to_llm_messages, to_unified_payload


def test_payload_basic(tmp_path: Path) -> None:
    path = tmp_path / "a.txt"
    path.write_text("hello", encoding="utf-8")

    payload = build_payload(path)
    assert str(payload.source).endswith("a.txt")
    assert payload.mime.startswith("text/")
    assert payload.text == "hello"
    assert payload.meta["size"] == 5
    assert payload.meta["line_count"] == 1
    assert payload.meta["word_count"] == 1

    unified = to_unified_payload(path)
    assert unified["doc_type"] == "text"
    assert unified["meta"]["word_count"] == 1
    assert unified["confidence"] > 0
    assert unified["validation"]["warnings"] == []


def test_payload_counts_ignore_trailing_newlines(tmp_path: Path) -> None:
    path = tmp_path / "multiline.txt"
    path.write_text("first line\nsecond line\n", encoding="utf-8")

    payload = build_payload(path)
    assert payload.meta["line_count"] == 2
    assert payload.meta["word_count"] == 4
    assert payload.meta["char_count"] == len("first line\nsecond line\n")

    unified = to_unified_payload(path)
    assert unified["meta"]["line_count"] == 2
    assert unified["meta"]["word_count"] == 4


def test_json_payload_includes_schema(tmp_path: Path) -> None:
    path = tmp_path / "data.json"
    path.write_text(json.dumps({"a": 1, "b": {"c": 2}}), encoding="utf-8")

    payload = build_payload(path)

    assert payload.mime == "application/json"
    assert payload.meta["json_valid"] is True
    assert payload.meta["json_top_level_type"] == "dict"
    assert payload.meta["json_top_level_keys"] == ["a", "b"]
    assert "word_count" not in payload.meta


def test_log_payload_summarizes_entries(tmp_path: Path) -> None:
    path = tmp_path / "events.log"
    path.write_text(
        "\n".join(
            [
                "2024-01-01 10:00:00 INFO service starting",
                "2024-01-01 10:01:00 WARN cache miss detected",
                "2024-01-01 10:02:00 ERROR critical failure",
                "    Traceback (most recent call last):",
                "    ValueError: boom",
            ]
        ),
        encoding="utf-8",
    )

    payload = build_payload(path)

    assert payload.meta["log_line_count"] == 3
    assert payload.meta["log_levels"] == ["ERROR", "WARN", "INFO"]
    assert payload.meta["log_first_timestamp"].startswith("2024-01-01T10:00:00")
    assert payload.meta["log_last_timestamp"].startswith("2024-01-01T10:02:00")
    assert payload.meta["log_level_counts"] == {"ERROR": 1, "WARN": 1, "INFO": 1}
    assert payload.meta["log_high_severity_count"] == 1
    assert payload.meta["log_has_multiline_entries"] is True
    assert payload.meta["log_examples"][0]["level"] == "INFO"
    assert "critical failure" in payload.meta["log_examples"][-1]["message"]

    unified = to_unified_payload(path)
    assert unified["doc_type"] == "log"
    assert any(block["type"] == "log" for block in unified["text_blocks"])
    assert "Log file contains high-severity entries." in unified["validation"]["warnings"]


def test_to_llm_messages_chunking(tmp_path: Path) -> None:
    path = tmp_path / "big.txt"
    path.write_text("a" * 5000, encoding="utf-8")

    payload = build_payload(path)
    messages = to_llm_messages(payload, chunk_size=2000)

    assert len(messages) == 3
    assert "part 1" in messages[0]["content"]
    assert "part 3" in messages[-1]["content"]


def test_html_payload_extracts_text(tmp_path: Path) -> None:
    html = """
    <html>
      <head><title>Sample Page</title></head>
      <body>
        <h1>Heading One</h1>
        <p>Hello <strong>world</strong>!</p>
      </body>
    </html>
    """
    path = tmp_path / "page.html"
    path.write_text(html, encoding="utf-8")

    payload = build_payload(path)

    assert "Hello" in payload.text and "world" in payload.text
    assert payload.meta["html_title"] == "Sample Page"
    assert "Heading One" in payload.meta["html_headings"]
    assert payload.meta["word_count"] >= 3

    unified = to_unified_payload(path)
    assert any(block["type"] == "header" for block in unified["text_blocks"])


def test_docx_payload_includes_tables(tmp_path: Path) -> None:
    path = tmp_path / "report.docx"
    document = DocxDocument()
    document.add_paragraph("Docx hello world")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    document.save(path)

    payload = build_payload(path)

    assert "Docx hello world" in payload.text
    assert payload.meta["docx_paragraphs"] >= 1
    assert payload.meta["docx_tables"] == 1
    assert payload.meta["word_count"] >= 3

    unified = to_unified_payload(path)
    assert any(table["rows"] for table in unified["tables"])


def test_xlsx_payload_renders_cells(tmp_path: Path) -> None:
    path = tmp_path / "data.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "SheetA"
    sheet.append(["Name", "Value"])
    sheet.append(["Alpha", 1])
    workbook.save(path)

    payload = build_payload(path)

    assert "# SheetA" in payload.text
    assert "Alpha" in payload.text
    assert payload.meta["workbook_sheet_count"] == 1
    assert payload.meta["workbook_cell_values"] == 4
    assert payload.meta["word_count"] >= 2

    unified = to_unified_payload(path)
    assert any(table["rows"] for table in unified["tables"])


def test_pdf_payload_reports_pages(tmp_path: Path) -> None:
    pdf_bytes = base64.b64decode(
        "JVBERi0xLjMKJcTl8uXrp/Og0MTGCjQgMCBvYmoKPDwgL0xlbmd0aCAxMTIgPj4Kc3RyZWFtCkJUIC9GMSAyNCBUZiA1MCAxNTAgVGQKKEhlbGxvIFBERik"
        "gVGoKRVQKZW5kc3RyZWFtCmVuZG9iago1IDAgb2JqCjw8IC9UeXBlIC9Gb250IC9TdWJ0eXBlIC9UeXBlMSAvTmFtZSAvRjEgL0Jhc2VGb250IC9IZWx2ZXRpY2EgPj4"
        "KZW5kb2JqCjMgMCBvYmoKPDwgL1R5cGUgL1BhZ2UgL1BhcmVudCAyIDAgUiAvTWVkaWFCb3ggWyAwIDAgMjAwIDIwMCBdIC9Db250ZW50cyA0IDAgUiAvUmVzb3VyY2V"
        "zIDw8IC9Gb250IDw8IC9GMSA1IDAgUiA+PiA+PiA+PgplbmRvYmoKMiAwIG9iago8PCAvVHlwZSAvUGFnZXMgL0tpZHMgWyAzIDAgUiBdIC9Db3VudCAxID4+CmVuZG9"
        "iagoxIDAgb2JqCjw8IC9UeXBlIC9DYXRhbG9nIC9QYWdlcyAyIDAgUiA+PgplbmRvYmoKeHJlZg0KMCA2DQowMDAwMDAwMDAwIDY1NTM1IGYgDQowMDAwMDAwMDEwIDA"
        "wMDAwIG4gDQowMDAwMDAwMDYwIDAwMDAwIG4gDQowMDAwMDAwMTE3IDAwMDAwIG4gDQowMDAwMDAwMjIwIDAwMDAwIG4gDQowMDAwMDAwMzAzIDAwMDAwIG4gDQp0cmF"
        "pbGVyDQo8PCAvU2l6ZSA2IC9Sb290IDEgMCBSID4+DQplbmR0cmFpbGVyDQpzdGFydHhyZWYNCjM0OQ0KJSVFT0YN"
    )
    path = tmp_path / "hello.pdf"
    path.write_bytes(pdf_bytes)

    payload = build_payload(path)

    assert "Hello PDF" in payload.text
    assert payload.meta["pdf_page_count"] == 1
    assert payload.meta["pdf_has_text"] is True
    assert payload.meta["word_count"] >= 2

    unified = to_unified_payload(path)
    assert any(block["type"] == "paragraph" for block in unified["text_blocks"])


def test_email_payload_includes_entities(tmp_path: Path) -> None:
    path = tmp_path / "message.eml"
    message = (
        "From: Sender <sender@example.com>\n"
        "To: receiver@example.com\n"
        "Subject: Invoice\n"
        "Date: Wed, 02 Aug 2023 12:00:00 +0000\n"
        "Content-Type: text/plain; charset=utf-8\n"
        "\n"
        "Hello team,\n"
        "Total due: $1,234.56\n"
    )
    path.write_text(message, encoding="utf-8")

    payload = build_payload(path)
    assert payload.meta["email_subject"] == "Invoice"
    assert payload.meta["email_to"] == ["receiver@example.com"]

    unified = to_unified_payload(path)
    assert unified["doc_type"] == "email"
    assert any(party["role"] == "from" for party in unified["parties"])
    assert any(amount["value"] == "1234.56" for amount in unified["amounts"])
    assert any(date["origin"] == "meta.email_date" for date in unified["dates"])


def test_zip_payload_surfaces_entries(tmp_path: Path) -> None:
    path = tmp_path / "bundle.zip"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("logs/app.log", "2024-01-01 started\n")
        archive.writestr("data.csv", "name,value\nalpha,10\n")
        archive.writestr("nested/archive.zip", b"PK\x03\x04")

    payload = build_payload(path)
    assert payload.meta["zip_entry_count"] == 3
    assert payload.meta["zip_contains_nested"] is True

    unified = to_unified_payload(path)
    assert unified["doc_type"] == "archive"
    assert any(att["name"] == "data.csv" for att in unified["attachments"])
    assert any(block["type"] == "attachment" for block in unified["text_blocks"])


def test_rtf_payload_roundtrip(tmp_path: Path) -> None:
    rtf = r"{\rtf1\ansi\deff0 {\fonttbl{\f0 Courier;}}\n\f0\fs20 Hello \\b World\\b0!}"
    path = tmp_path / "note.rtf"
    path.write_text(rtf, encoding="utf-8")

    payload = build_payload(path)
    assert "Hello" in payload.text
    assert "World" in payload.text
    assert payload.meta.get("rtf_decoder") in {"striprtf", "fallback"}

    unified = to_unified_payload(path)
    assert unified["doc_type"] == "text"
    assert unified["highlights"]["summary"]


def test_pptx_payload_extracts_slides(tmp_path: Path) -> None:
    path = tmp_path / "deck.pptx"
    slide_xml = """<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>
    <p:sld xmlns:p=\"http://schemas.openxmlformats.org/presentationml/2006/main\"
           xmlns:a=\"http://schemas.openxmlformats.org/drawingml/2006/main\"
           xmlns:r=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\">
      <p:cSld>
        <p:spTree>
          <p:sp>
            <p:txBody>
              <a:bodyPr/>
              <a:lstStyle/>
              <a:p>
                <a:r><a:t>Deck Title</a:t></a:r>
              </a:p>
            </p:txBody>
          </p:sp>
          <p:sp>
            <p:txBody>
              <a:bodyPr/>
              <a:lstStyle/>
              <a:p>
                <a:r><a:t>First bullet</a:t></a:r>
              </a:p>
            </p:txBody>
          </p:sp>
        </p:spTree>
      </p:cSld>
    </p:sld>
    """
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("ppt/slides/slide1.xml", slide_xml)

    payload = build_payload(path)
    assert "Deck Title" in payload.text
    assert payload.meta["pptx_slide_count"] == 1
    assert payload.meta["pptx_slide_titles"] == ["Deck Title"]

    unified = to_unified_payload(path)
    assert unified["doc_type"] == "presentation"
    assert unified["llm_ready"]["markdown"].startswith("# Document: Presentation")
    assert unified["highlights"]["summary"]
    assert unified["llm_ready"]["chunks"]
