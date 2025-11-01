# SPDX-License-Identifier: AGPL-3.0-or-later
from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path
from typing import Dict

import pytest

from docx import Document as DocxDocument
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont, PngImagePlugin
from pypdf import PdfWriter

from sr_adapter.pipeline import batch_convert, convert, stream_convert
from sr_adapter.settings import reset_settings_cache
from sr_adapter.profiles import LLMPolicy, ProcessingProfile
from sr_adapter.parsers import parse_image, parse_ini, parse_txt
from sr_adapter.recipe import apply_recipe
from sr_adapter.schema import Block
from sr_adapter.sniff import detect_type


def _create_pdf(path: Path) -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with path.open("wb") as handle:
        writer.write(handle)


def _create_png(path: Path, text: str) -> None:
    image = Image.new("RGB", (320, 120), color="white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.load_default()
    except Exception:
        font = None
    draw.text((10, 40), text, fill="black", font=font)

    info = PngImagePlugin.PngInfo()
    info.add_text("Description", text)
    image.save(str(path), pnginfo=info)


def _create_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("Example Document", level=1)
    doc.add_paragraph("Body text")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Key"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "1"
    doc.save(str(path))


def _create_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Key"
    ws["B1"] = "Value"
    ws["A2"] = "Temp"
    ws["B2"] = 42
    wb.save(str(path))


def test_detect_type_handles_various_formats(tmp_path: Path) -> None:
    txt = tmp_path / "sample.txt"
    txt.write_text("plain", encoding="utf-8")

    md = tmp_path / "sample.md"
    md.write_text("# Title\n- item", encoding="utf-8")

    pdf = tmp_path / "sample.pdf"
    _create_pdf(pdf)

    docx = tmp_path / "sample.docx"
    _create_docx(docx)

    xlsx = tmp_path / "sample.xlsx"
    _create_xlsx(xlsx)

    yaml_file = tmp_path / "config.yaml"
    yaml_file.write_text("service: demo", encoding="utf-8")

    jsonl = tmp_path / "records.jsonl"
    jsonl.write_text("{}\n", encoding="utf-8")

    toml = tmp_path / "config.toml"
    toml.write_text("title = \"Example\"\n", encoding="utf-8")

    cfg = tmp_path / "settings.cfg"
    cfg.write_text("host localhost\n", encoding="utf-8")

    props = tmp_path / "application.properties"
    props.write_text("user => demo\n", encoding="utf-8")

    png = tmp_path / "invoice.png"
    _create_png(png, "Invoice #42 Total 19.99")

    assert detect_type(txt) == "text"
    assert detect_type(md) == "md"
    assert detect_type(pdf) == "pdf"
    assert detect_type(docx) == "docx"
    assert detect_type(xlsx) == "xlsx"
    assert detect_type(yaml_file) == "yaml"
    assert detect_type(jsonl) == "jsonl"
    assert detect_type(toml) == "toml"
    assert detect_type(cfg) == "ini"
    assert detect_type(props) == "ini"
    assert detect_type(png) == "image"


def test_convert_with_recipe_applies_patterns(tmp_path: Path) -> None:
    log = tmp_path / "call.log"
    log.write_text(
        "[2024-01-01 09:30] Call started\nCaller: Alice\n- greeted\n",
        encoding="utf-8",
    )

    document = convert(log, recipe="call_log")
    assert document.meta["type"] == "text"
    assert len(document.blocks) >= 3
    assert document.blocks[0].type == "meta"
    assert document.blocks[1].type == "header"
    assert any(block.type == "list" for block in document.blocks)


def test_parse_txt_extracts_kv_and_log(tmp_path: Path) -> None:
    source = tmp_path / "structured.txt"
    source.write_text(
        "\n".join(
            [
                "Subject: Update",
                "Name -> Alice",
                "2024-01-01 09:30 status ok",
                "- bullet item",
            ]
        ),
        encoding="utf-8",
    )

    blocks = parse_txt(source)

    types = {block.type for block in blocks}
    assert "kv" in types
    assert "log" in types
    log_block = next(block for block in blocks if block.type == "log")
    assert log_block.attrs["timestamp"].startswith("2024-01-01")
    assert "status ok" in log_block.attrs["message"]
    name_block = next(block for block in blocks if block.type == "kv" and block.attrs.get("key") == "Name")
    assert name_block.attrs["value"] == "Alice"


def test_parse_txt_detects_pipe_table(tmp_path: Path) -> None:
    source = tmp_path / "table.txt"
    source.write_text(
        "\n".join(
            [
                "| Name | Age | City |",
                "| --- | --- | --- |",
                "| Alice | 30 | Tokyo |",
                "| Bob | 41 | Osaka |",
            ]
        ),
        encoding="utf-8",
    )

    blocks = parse_txt(source)

    table_block = next(block for block in blocks if block.type == "table")
    rows = json.loads(table_block.attrs["rows"])
    assert rows[0] == ["Name", "Age", "City"]
    assert rows[1][0] == "Alice"
    assert table_block.attrs["delimiter"] == "pipe"
    assert table_block.attrs["row_count"] == 3


def test_parse_txt_detects_whitespace_table(tmp_path: Path) -> None:
    source = tmp_path / "whitespace_table.txt"
    source.write_text(
        "\n".join(
            [
                "Name    Score    Notes",
                "Alice   98       OK",
                "Bob     87       Needs follow-up",
            ]
        ),
        encoding="utf-8",
    )

    blocks = parse_txt(source)

    table_block = next(block for block in blocks if block.type == "table")
    rows = json.loads(table_block.attrs["rows"])
    assert len(rows) == 3
    assert rows[2][2] == "Needs follow-up"
    assert table_block.attrs["delimiter"] == "whitespace"


def _write_documents(tmp_path: Path, count: int = 3) -> list[Path]:
    paths: list[Path] = []
    for idx in range(count):
        path = tmp_path / f"doc-{idx}.txt"
        path.write_text(f"Document {idx}", encoding="utf-8")
        paths.append(path)
    return paths


def test_batch_convert_thread_backend(tmp_path: Path, monkeypatch) -> None:
    files = _write_documents(tmp_path, count=2)

    calls: dict[str, object] = {}

    def _fake_threadpool(func, items, *, workers):
        calls["workers"] = workers
        return [func(item) for item in items]

    monkeypatch.setattr("sr_adapter.pipeline.run_threadpool", _fake_threadpool)
    reset_settings_cache()
    documents = batch_convert(files, recipe="default", backend="threadpool", concurrency=3)

    assert len(documents) == 2
    assert calls["workers"] == 3


def test_batch_convert_async_backend(tmp_path: Path, monkeypatch) -> None:
    files = _write_documents(tmp_path, count=2)

    def _fake_asyncio(func, items, *, workers):
        return [func(item) for item in items]

    monkeypatch.setattr("sr_adapter.pipeline.run_asyncio", _fake_asyncio)
    reset_settings_cache()
    documents = batch_convert(files, recipe="default", backend="asyncio", concurrency=0)

    assert len(documents) == 2


def test_batch_convert_dask_backend(tmp_path: Path, monkeypatch) -> None:
    files = _write_documents(tmp_path, count=2)

    captured: dict[str, object] = {}

    def _fake_dask(func, items, *, scheduler, workers):
        captured["scheduler"] = scheduler
        captured["workers"] = workers
        return [func(item) for item in items]

    monkeypatch.setattr("sr_adapter.pipeline.run_dask", _fake_dask)
    reset_settings_cache()
    documents = batch_convert(
        files,
        recipe="default",
        backend="dask",
        dask_scheduler="threads",
        concurrency=5,
    )

    assert len(documents) == 2
    assert captured["scheduler"] == "threads"
    assert captured["workers"] == 5


def test_batch_convert_ray_backend(tmp_path: Path, monkeypatch) -> None:
    files = _write_documents(tmp_path, count=2)

    captured: dict[str, object] = {}

    def _fake_ray(func, items, *, address, workers):
        captured["address"] = address
        captured["workers"] = workers
        return [func(item) for item in items]

    monkeypatch.setattr("sr_adapter.pipeline.run_ray", _fake_ray)
    reset_settings_cache()
    documents = batch_convert(
        files,
        recipe="default",
        backend="ray",
        ray_address="auto",
        concurrency=0,
    )

    assert len(documents) == 2
    assert captured["address"] == "auto"
    # With concurrency=0 the default max_workers from settings.yaml is used (4)
    assert captured["workers"] == 4


def test_parse_ini_coerces_space_pairs(tmp_path: Path) -> None:
    source = tmp_path / "messy.cfg"
    source.write_text(
        "\n".join(
            [
                "# comment",
                "host localhost",
                "port 5432",
                "",
                "[service]",
                "enabled true",
                "timeout 30",
                "path => /var/lib/app",
            ]
        ),
        encoding="utf-8",
    )

    blocks = parse_ini(source)

    root = next(block for block in blocks if block.attrs.get("key") == "<root>")
    assert root.attrs["section_count"] == 1
    assert "service" in root.attrs["sections"]
    assert "host" in root.attrs.get("default_keys", [])
    assert root.attrs["coerced_pairs"] == 5

    host = next(block for block in blocks if block.attrs.get("key") == "<defaults>.host")
    assert host.attrs["value"] == "localhost"

    service_path = next(block for block in blocks if block.attrs.get("key") == "service.path")
    assert service_path.attrs["value"] == "/var/lib/app"


def test_convert_ini_uses_structured_parser(tmp_path: Path) -> None:
    source = tmp_path / "app.properties"
    source.write_text(
        "\n".join(
            [
                "name demo",
                "enabled true",
                "[limits]",
                "max 10",
            ]
        ),
        encoding="utf-8",
    )

    document = convert(source, recipe="default")

    assert document.meta["type"] == "ini"
    head = document.blocks[0]
    assert head.attrs.get("coerced_pairs", 0) >= 3
    assert any(block.attrs.get("key") == "<defaults>.name" for block in document.blocks if block.type == "kv")
    assert any(block.attrs.get("key") == "limits.max" for block in document.blocks if block.type == "kv")


def test_convert_image_uses_metadata_text(tmp_path: Path) -> None:
    image = tmp_path / "invoice.png"
    _create_png(image, "Invoice #42 Total 19.99")

    document = convert(image, recipe="default")

    assert document.meta["type"] == "image"
    joined = "\n".join(block.text for block in document.blocks)
    assert "Invoice #42" in joined
    meta_blocks = [block for block in document.blocks if block.type == "metadata"]
    assert meta_blocks
    assert any(block.attrs.get("image_source") == "metadata" for block in meta_blocks)
    assert any(block.attrs.get("image_has_text") is True for block in meta_blocks)


def test_parse_image_propagates_language_metadata(monkeypatch, tmp_path: Path) -> None:
    image = tmp_path / "ocr.png"
    _create_png(image, "hello")

    fake_meta = {"image_has_text": True, "image_ocr_languages": ["eng", "jpn"]}
    fake_segments = [
        {"text": "hello", "source": "ocr", "kind": "ocr", "confidence": 0.9, "order": 1},
    ]

    def fake_extract(path: Path):
        assert Path(path) == image
        return "hello", dict(fake_meta), list(fake_segments)

    monkeypatch.setattr("sr_adapter.parsers._extract_image_text", fake_extract)

    blocks = parse_image(image)

    summary = next(block for block in blocks if block.type == "metadata" and block.text == "Image summary")
    assert summary.attrs.get("image_ocr_languages") == ["eng", "jpn"]

    ocr_blocks = [block for block in blocks if block.attrs.get("image_source") == "ocr"]
    assert ocr_blocks
    assert any(block.attrs.get("ocr_languages") == ["eng", "jpn"] for block in ocr_blocks)


def test_parse_image_orders_segments_with_native_kernel(monkeypatch, tmp_path: Path) -> None:
    image = tmp_path / "layout.png"
    _create_png(image, "placeholder")

    fake_meta: Dict[str, object] = {"image_has_text": True}
    fake_segments = [
        {"text": "second line", "source": "ocr", "kind": "ocr", "confidence": 0.45, "bbox": (0, 60, 120, 90), "order": 2},
        {"text": "first line", "source": "ocr", "kind": "ocr", "confidence": 0.8, "bbox": (0, 10, 120, 35), "order": 1},
    ]

    def fake_extract(path: Path):
        assert Path(path) == image
        return "", dict(fake_meta), list(fake_segments)

    monkeypatch.setattr("sr_adapter.parsers._extract_image_text", fake_extract)

    blocks = parse_image(image)

    ordered = [block for block in blocks if block.type != "metadata"]
    assert ordered
    assert ordered[0].text == "first line"
    assert ordered[0].attrs.get("layout_kernel") == "native-cpp-v1"
    assert "layout_calibrated_threshold" in ordered[0].attrs


def test_recipe_fallback_preserves_attrs(tmp_path: Path) -> None:
    block = Block(type="paragraph", text="unknown content", attrs={"a": "1"})
    processed = apply_recipe([block], "default")
    assert processed[0].type == "paragraph"
    assert processed[0].attrs["a"] == "1"


def test_stream_convert_matches_convert(tmp_path: Path) -> None:
    source = tmp_path / "note.txt"
    source.write_text("Alpha\nBeta\n", encoding="utf-8")

    streamed = list(stream_convert(source, recipe="default"))
    document = convert(source, recipe="default")

    assert [block.text for block in streamed] == [block.text for block in document.blocks]

    limited = list(stream_convert(source, recipe="default", max_blocks=1))
    assert len(limited) == 1


def test_cli_convert_produces_jsonl(tmp_path: Path) -> None:
    source = tmp_path / "sensor.log"
    source.write_text(
        "Sensor: A1\n2024-01-01T00:00:00\nTemp: 21.5\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.jsonl"

    result = subprocess.run(
        [sys.executable, "-m", "sr_adapter.cli", "convert", str(source), "--recipe", "sensor_log", "--out", str(out)],
        check=True,
        capture_output=True,
        text=True,
    )
    assert result.returncode == 0
    lines = out.read_text(encoding="utf-8").splitlines()
    assert len(lines) == 1
    payload = json.loads(lines[0])
    assert payload["meta"]["type"] == "text"
    assert any(block["type"] == "kv" for block in payload["blocks"])


def test_convert_eml_extracts_headers(tmp_path: Path) -> None:
    eml = tmp_path / "mail.eml"
    eml.write_text(
        "\n".join(
            [
                "From: Example <sender@example.com>",
                "To: Recipient <user@example.com>",
                "Subject: Hello",
                "MIME-Version: 1.0",
                "Content-Type: text/plain; charset=utf-8",
                "",
                "Hello world",
            ]
        ),
        encoding="utf-8",
    )

    document = convert(eml, recipe="default")

    assert document.meta["type"] == "eml"
    assert document.blocks[0].type == "meta"
    assert document.blocks[0].attrs["subject"] == "Hello"
    assert any(block.type in {"paragraph", "kv"} for block in document.blocks[1:])


def test_convert_ics_yields_event_blocks(tmp_path: Path) -> None:
    ics = tmp_path / "event.ics"
    ics.write_text(
        "\n".join(
            [
                "BEGIN:VCALENDAR",
                "VERSION:2.0",
                "BEGIN:VEVENT",
                "SUMMARY:Sync",
                "DTSTART:20240201T120000Z",
                "DTEND:20240201T123000Z",
                "LOCATION:Video",
                "DESCRIPTION:Weekly sync",
                "END:VEVENT",
                "END:VCALENDAR",
            ]
        ),
        encoding="utf-8",
    )

    document = convert(ics, recipe="default")

    assert document.meta["type"] == "ics"
    assert any(block.type == "event" for block in document.blocks)
    event = next(block for block in document.blocks if block.type == "event")
    assert event.attrs["summary"] == "Sync"


def test_convert_json_extracts_nested_keys(tmp_path: Path) -> None:
    payload = {
        "service": {"host": "localhost", "ports": [80, 443]},
        "debug": True,
    }
    src = tmp_path / "config.json"
    src.write_text(json.dumps(payload), encoding="utf-8")

    document = convert(src, recipe="default")

    assert document.meta["type"] == "json"
    kv_keys = {block.attrs.get("key") for block in document.blocks if block.type == "kv"}
    assert "service.host" in kv_keys
    assert "service.ports[0]" in kv_keys
    assert any(block.attrs.get("type") == "array" for block in document.blocks if block.type == "metadata")
    host_block = next(block for block in document.blocks if block.attrs.get("key") == "service.host")
    assert host_block.attrs["path_r"] == ".data$service$host"
    assert host_block.attrs["path_glue"] == "{.data$service$host}"
    port_block = next(block for block in document.blocks if block.attrs.get("key") == "service.ports[0]")
    assert port_block.attrs["path_r"] == ".data$service$ports[[1]]"
    assert port_block.attrs["path_glue"] == "{.data$service$ports[[1]]}"
    service_meta = next(block for block in document.blocks if block.attrs.get("key") == "service")
    assert service_meta.attrs.get("sample_values") == ["localhost"]
    ports_meta = next(block for block in document.blocks if block.attrs.get("key") == "service.ports")
    assert ports_meta.attrs.get("sample_values") == ["80", "443"]


def test_convert_json_repairs_loose_json(tmp_path: Path) -> None:
    src = tmp_path / "messy.json"
    src.write_text(
        "\n".join(
            [
                "{",
                "  'service': {host: 'localhost', ports: [80, 443,],},",
                "  // legacy flags",
                "  'flags': ['alpha', 'beta',],",
                "}",
            ]
        ),
        encoding="utf-8",
    )

    document = convert(src, recipe="default")

    assert document.meta["type"] == "json"
    host_block = next(block for block in document.blocks if block.type == "kv" and block.attrs.get("key") == "service.host")
    assert host_block.attrs["value"] == "localhost"
    root_meta = next(block for block in document.blocks if block.attrs.get("key") == "<root>")
    assert root_meta.attrs.get("coerced_from") in {"sanitized", "yaml"}


def test_convert_yaml_emits_summary_and_kv(tmp_path: Path) -> None:
    src = tmp_path / "config.yaml"
    src.write_text(
        "\n".join([
            "service:",
            "  host: localhost",
            "  retries: 3",
            "  tags:",
            "    - api",
            "    - v1",
        ]),
        encoding="utf-8",
    )

    document = convert(src, recipe="default")

    assert document.meta["type"] == "yaml"
    kv_keys = {block.attrs.get("key") for block in document.blocks if block.type == "kv"}
    assert "service.host" in kv_keys
    summary_block = next(block for block in document.blocks if block.text.startswith("YAML:"))
    assert summary_block.attrs["path_r"] == ".data"
    assert summary_block.attrs["path_glue"] == "{.data}"
    assert summary_block.attrs["key"] == "<root>"


def test_convert_toml_handles_tables(tmp_path: Path) -> None:
    src = tmp_path / "settings.toml"
    src.write_text(
        "\n".join([
            "[database]",
            "user = \"admin\"",
            "[service]",
            "enabled = true",
        ]),
        encoding="utf-8",
    )

    document = convert(src, recipe="default")

    assert document.meta["type"] == "toml"
    kv_keys = {block.attrs.get("key") for block in document.blocks if block.type == "kv"}
    assert "database.user" in kv_keys
    assert any(block.attrs.get("type") == "object" for block in document.blocks if block.type == "metadata")


def test_convert_jsonl_breaks_into_records(tmp_path: Path) -> None:
    src = tmp_path / "records.jsonl"
    src.write_text(
        "\n".join([
            json.dumps({"id": 1, "value": "alpha"}),
            json.dumps({"id": 2, "value": "beta"}),
        ]),
        encoding="utf-8",
    )

    document = convert(src, recipe="default")

    assert document.meta["type"] == "jsonl"
    kv_keys = {block.attrs.get("key") for block in document.blocks if block.type == "kv"}
    assert "record[1].id" in kv_keys
    assert any(block.attrs.get("type") == "object" for block in document.blocks if block.attrs.get("key") == "record[1]")
    record_summary = next(block for block in document.blocks if block.attrs.get("key") == "record[1]")
    assert record_summary.attrs["path_r"] == ".data$record[[1]]"
    assert record_summary.attrs["path_glue"] == "{.data$record[[1]]}"
    record_id = next(block for block in document.blocks if block.attrs.get("key") == "record[1].id")
    assert record_id.attrs["path_r"] == ".data$record[[1]]$id"
    assert record_id.attrs["path_glue"] == "{.data$record[[1]]$id}"


def test_parse_txt_refines_large_paragraphs(tmp_path: Path) -> None:
    body = " ".join(["Sentence number %d." % i for i in range(1, 80)])
    path = tmp_path / "long.txt"
    path.write_text(body, encoding="utf-8")

    blocks = parse_txt(path)

    assert len(blocks) > 1
    assert all(len(block.text) <= 600 for block in blocks)


def test_convert_profile_forwards_llm_policy(monkeypatch, tmp_path: Path) -> None:
    path = tmp_path / "profile.txt"
    path.write_text("alpha\nbeta", encoding="utf-8")

    captured: dict = {}

    def _fake_escalate(blocks, recipe_name, **kwargs):
        captured.update(kwargs)
        return list(blocks)

    monkeypatch.setattr("sr_adapter.pipeline.escalate_low_conf", _fake_escalate)

    policy = LLMPolicy(max_confidence=1.0, limit_block_types=("paragraph",), max_blocks=2)
    profile = ProcessingProfile(name="test", llm_policy=policy, warm_runtime=False)

    document = convert(path, recipe="default", profile=profile)

    assert captured["max_confidence"] == 1.0
    assert captured["allow_types"] == ("paragraph",)
    assert captured["limit"] == 2
    assert document.meta["processing_profile"] == "test"
    assert document.meta["llm_policy"]["max_confidence"] == pytest.approx(1.0)

