# SPDX-License-Identifier: AGPL-3.0-or-later
"""Collection of lightweight parsers that emit :class:`Block` objects."""

from __future__ import annotations

import configparser
import csv
import json
import re
import zipfile
import xml.etree.ElementTree as ET
from collections import Counter
from collections.abc import Mapping, Sequence
from dataclasses import dataclass
from email import policy
from email.parser import BytesParser
from functools import lru_cache
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Sequence, Tuple

import yaml
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from .loaders import _extract_image_text
from .schema import BBox, Block, Provenance, clone_model
from .visual import LayoutCandidate, VisualLayoutAnalyzer


_LIST_MARKERS = re.compile(r"^(?:[-*\u2022\u30fb]|\d+[.)])\s+")
_TIMESTAMP_PREFIX = re.compile(r"^\[?\d{4}[-/]\d{2}[-/]\d{2}")
_SENTENCE_BOUNDARY = re.compile(r"(?<=[.!?。！？])\s+(?=[\w\"'(])")
_KV_PATTERN = re.compile(
    r"^(?P<key>[\w .#/@&()'\-]{1,64})\s*(?:=>|->|[:=]|：)\s*(?P<value>.+)$"
)
_LOG_LINE = re.compile(
    r"^\[?(?P<ts>\d{4}[-/]\d{2}[-/]\d{2}(?:[ T]\d{2}:\d{2}(?::\d{2})?(?:Z|[+-]\d{2}:?\d{2})?)?)\]?\s*(?P<body>.*)$"
)
_MAX_CHARS_PER_CHUNK = 600
_STRUCTURED_BLOCK_LIMIT = 400
_R_IDENTIFIER = re.compile(r"^[A-Za-z.][A-Za-z0-9._]*$")
_JSON_COMMENT_RE = re.compile(r"//.*?$|/\*.*?\*/", re.DOTALL | re.MULTILINE)
_TRAILING_COMMA_RE = re.compile(r",(\s*[}\]])")
_INI_ARROW_RE = re.compile(
    r"^(?P<indent>\s*)(?P<key>[^\s:=#;\[\]][^:=#;]*?)\s*(?:=>|->)\s*(?P<value>.+)$"
)
_WHITESPACE_TABLE_SPLIT = re.compile(r"\s{2,}")
_MARKDOWN_DIVIDER = re.compile(r"^:?-{3,}:?$")


@dataclass(frozen=True)
class _PathSegment:
    kind: str  # "key" or "index"
    value: str | int
    base: int | None = None  # index display base when kind == "index"


PathTuple = tuple[_PathSegment, ...]


def _new_block(block_type: str, text: str, source: Path, confidence: float = 0.5) -> Block:
    return Block(type=block_type, text=text, source=str(source), confidence=confidence)


def _split_paragraphs(text: str) -> Iterable[str]:
    buf: List[str] = []
    for line in text.splitlines():
        if line.strip():
            buf.append(line)
        elif buf:
            yield "\n".join(buf)
            buf.clear()
    if buf:
        yield "\n".join(buf)


def _classify_chunk(chunk: str) -> str:
    stripped = chunk.strip()
    if not stripped:
        return "other"
    if stripped.startswith("#") and stripped.count("#") <= 6:
        return "heading"
    if stripped.upper() == stripped and len(stripped.split()) <= 6:
        return "heading"
    if _LOG_LINE.match(stripped):
        return "log"
    if _KV_PATTERN.match(stripped):
        return "kv"
    lines = stripped.splitlines()
    if len(lines) == 1 and len(stripped) <= 80:
        lowered = stripped.lower()
        if any(lowered.startswith(prefix) for prefix in ("title:", "subject:", "from:", "to:", "cc:")):
            return "kv"
    if all(_LIST_MARKERS.match(line) for line in lines):
        return "list_item"
    if _TIMESTAMP_PREFIX.match(stripped):
        return "metadata"
    if all(":" in line for line in lines) and len(lines) <= 6:
        return "metadata"
    return "paragraph"


def _block_from_chunk(chunk: str, source: Path) -> Block:
    chunk = chunk.strip()
    block_type = _classify_chunk(chunk)
    if block_type == "kv":
        match = _KV_PATTERN.match(chunk)
        attrs: Dict[str, str] = {}
        if match:
            attrs = {"key": match.group("key").strip(), "value": match.group("value").strip()}
        return Block(
            type="kv",
            text=chunk,
            attrs=attrs,
            source=str(source),
            confidence=0.75,
        )
    if block_type == "log":
        match = _LOG_LINE.match(chunk)
        attrs = {}
        if match:
            attrs = {
                "timestamp": match.group("ts"),
                "message": match.group("body").strip(),
            }
        return Block(
            type="log",
            text=chunk,
            attrs=attrs,
            source=str(source),
            confidence=0.6,
        )
    if block_type == "list_item":
        return _new_block("list_item", chunk, source, confidence=0.7)
    if block_type == "heading":
        return _new_block("heading", chunk, source, confidence=0.8)
    if block_type == "metadata":
        return _new_block("metadata", chunk, source, confidence=0.6)
    if block_type == "other":
        return _new_block("other", chunk, source, confidence=0.3)
    return _new_block(block_type, chunk, source)


def _explode_structured_chunk(chunk: str) -> List[str]:
    if "\n" not in chunk:
        return [chunk]
    lines = [line for line in chunk.splitlines() if line.strip()]
    # Treat short structured segments (logs, bullet lists) line-by-line.
    if lines and all(len(line) < 200 for line in lines) and len(lines) <= 16:
        return lines
    return [chunk]


def _shatter_chunk(chunk: str) -> List[str]:
    chunk = chunk.strip()
    if len(chunk) <= _MAX_CHARS_PER_CHUNK:
        return [chunk]
    sentences = [seg.strip() for seg in _SENTENCE_BOUNDARY.split(chunk) if seg.strip()]
    if len(sentences) > 1:
        grouped: List[str] = []
        current = ""
        for sentence in sentences:
            if not current:
                current = sentence
                continue
            tentative = f"{current} {sentence}" if current else sentence
            if len(tentative) <= _MAX_CHARS_PER_CHUNK:
                current = tentative
            else:
                grouped.append(current)
                current = sentence
        if current:
            grouped.append(current)
        return grouped
    return [chunk[i : i + _MAX_CHARS_PER_CHUNK] for i in range(0, len(chunk), _MAX_CHARS_PER_CHUNK)]


def _split_with_delimiter(line: str, delimiter: str) -> List[str]:
    try:
        reader = csv.reader([line], delimiter=delimiter)
        row = next(reader, [])
    except Exception:
        return []
    cells = [cell.strip() for cell in row]
    if delimiter == "|":
        while cells and not cells[0]:
            cells.pop(0)
        while cells and not cells[-1]:
            cells.pop()
    return cells


def _split_with_whitespace(line: str) -> List[str]:
    parts = [segment.strip() for segment in _WHITESPACE_TABLE_SPLIT.split(line.strip())]
    return [part for part in parts if part]


def _coerce_tabular_rows(lines: Sequence[str]) -> tuple[str, List[List[str]]] | None:
    cleaned = [line for line in lines if line.strip()]
    if len(cleaned) < 2:
        return None

    candidates = [
        ("comma", lambda line: _split_with_delimiter(line, ",")),
        ("tab", lambda line: _split_with_delimiter(line, "\t")),
        ("semicolon", lambda line: _split_with_delimiter(line, ";")),
        ("pipe", lambda line: _split_with_delimiter(line, "|")),
        ("whitespace", _split_with_whitespace),
    ]

    best_label = ""
    best_rows: List[List[str]] = []
    best_score: tuple[int, int] | None = None

    for label, splitter in candidates:
        rows = [splitter(line) for line in cleaned]
        usable = [row for row in rows if len(row) >= 2]
        if len(usable) < 2:
            continue

        width_counts = Counter(len(row) for row in usable)
        top_width, freq = width_counts.most_common(1)[0]
        consistent = [row[:top_width] for row in usable if len(row) >= top_width]

        filtered = [
            row
            for row in consistent
            if not all(_MARKDOWN_DIVIDER.match(cell) for cell in row)
        ]
        if len(filtered) < 2:
            continue

        coverage = len(filtered)
        if coverage < max(2, len(cleaned) - 1):
            continue

        score = (coverage, top_width)
        if not best_score or score > best_score:
            best_score = score
            best_label = label
            best_rows = filtered

    if not best_rows:
        return None

    width = max(len(row) for row in best_rows)
    normalized = [row[:width] + [""] * (width - len(row)) for row in best_rows]
    return best_label, normalized


def _iter_refined_chunks(text: str) -> Iterable[str]:
    for chunk in _split_paragraphs(text):
        lines = [line for line in chunk.splitlines() if line.strip()]
        table_hint = _coerce_tabular_rows(lines) if len(lines) >= 2 else None
        pieces = _explode_structured_chunk(chunk) if table_hint is None else [chunk]
        for piece in pieces:
            for refined in _shatter_chunk(piece):
                cleaned = refined.strip()
                if cleaned:
                    yield cleaned


def _format_dot_path(path: PathTuple) -> str:
    parts: List[str] = []
    for segment in path:
        if segment.kind == "key":
            if parts:
                parts.append(".")
            parts.append(str(segment.value))
        elif segment.kind == "index":
            parts.append(f"[{segment.value}]")
    return "".join(parts)


def _format_label(path: PathTuple) -> str:
    dotted = _format_dot_path(path)
    return dotted or "<root>"


def _format_r_key(name: str) -> str:
    if _R_IDENTIFIER.match(name):
        return f"${name}"
    escaped = name.replace("\\", "\\\\").replace('"', '\\"')
    return f'[["{escaped}"]]'


@lru_cache(maxsize=4096)
def _render_r_path(path: PathTuple) -> str:
    cursor = ".data"
    for segment in path:
        if segment.kind == "key":
            cursor += _format_r_key(str(segment.value))
        elif segment.kind == "index":
            base = 0 if segment.base is None else segment.base
            index_value = int(segment.value) - base + 1
            cursor += f"[[{index_value}]]"
    return cursor


def _format_r_path(path: PathTuple) -> str:
    return _render_r_path(path)


@lru_cache(maxsize=4096)
def _render_r_tokens(path: PathTuple) -> Tuple[str, ...]:
    tokens = [".data"]
    for segment in path:
        if segment.kind == "key":
            tokens.append(_format_r_key(str(segment.value)))
        elif segment.kind == "index":
            base = 0 if segment.base is None else segment.base
            index_value = int(segment.value) - base + 1
            tokens.append(f"[[{index_value}]]")
    return tuple(tokens)


def _format_r_tokens(path: PathTuple) -> List[str]:
    return list(_render_r_tokens(path))


def _format_glue_path(path: PathTuple) -> str:
    return f"{{{_format_r_path(path)}}}"


def _extend_key(path: PathTuple, key: object) -> PathTuple:
    return path + (_PathSegment("key", str(key)),)


def _extend_index(path: PathTuple, index: int, *, base: int = 0) -> PathTuple:
    return path + (_PathSegment("index", index, base),)


def _stringify_scalar(value: object) -> str:
    if isinstance(value, str):
        return value
    if value is None:
        return "null"
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, (int, float)):
        return str(value)
    try:
        return json.dumps(value, ensure_ascii=False)
    except TypeError:
        return str(value)


def _strip_json_comments(text: str) -> str:
    def _replace(match: re.Match[str]) -> str:
        span = match.group(0)
        # Preserve newlines to keep downstream line numbers vaguely aligned.
        return "".join("\n" if ch == "\n" else " " for ch in span)

    return _JSON_COMMENT_RE.sub(_replace, text)


def _sanitize_json_like(text: str) -> str:
    cleaned = text.lstrip("\ufeff")
    cleaned = _strip_json_comments(cleaned)
    cleaned = _TRAILING_COMMA_RE.sub(r"\1", cleaned)
    return cleaned


def _load_json_like(raw: str) -> tuple[object, str | None]:
    primary = raw.lstrip("\ufeff")
    try:
        return json.loads(primary), None
    except json.JSONDecodeError:
        pass

    sanitized = _sanitize_json_like(primary)
    if sanitized != primary:
        try:
            return json.loads(sanitized), "sanitized"
        except json.JSONDecodeError:
            pass

    candidate = sanitized.strip()
    if candidate:
        try:
            data = yaml.safe_load(sanitized)
        except yaml.YAMLError:
            data = None
        else:
            return data, "yaml"

    raise ValueError("Unable to coerce JSON input")


def _prepare_ini_input(text: str) -> tuple[str, dict[str, object], bool, bool]:
    cleaned = text.lstrip("\ufeff")
    space_converted = 0
    arrow_converted = 0
    leading_pairs = False
    seen_section = False
    lines: List[str] = []

    for raw_line in cleaned.splitlines():
        stripped = raw_line.strip()
        if not stripped or stripped.startswith(("#", ";")):
            lines.append(raw_line)
            continue
        if stripped.startswith("["):
            lines.append(raw_line)
            seen_section = True
            continue

        line = raw_line
        line_is_pair = False

        arrow_match = _INI_ARROW_RE.match(raw_line)
        if arrow_match:
            indent = arrow_match.group("indent")
            key = arrow_match.group("key").strip()
            value = arrow_match.group("value").strip()
            line = f"{indent}{key} = {value}"
            arrow_converted += 1
            line_is_pair = True
        elif "=" in raw_line or ":" in raw_line:
            line_is_pair = True
        else:
            parts = stripped.split(None, 1)
            if len(parts) == 2 and not any(ch in parts[0] for ch in "[]=#:;"):
                indent = raw_line[: len(raw_line) - len(raw_line.lstrip())]
                key, value = parts
                line = f"{indent}{key} = {value}"
                space_converted += 1
                line_is_pair = True

        lines.append(line)

        if line_is_pair and not seen_section:
            leading_pairs = True

    sanitized = "\n".join(lines)
    has_section = seen_section or any(line.lstrip().startswith("[") for line in lines)
    meta: dict[str, object] = {}
    if arrow_converted:
        meta["coerced_arrow_pairs"] = arrow_converted
    if space_converted:
        meta["coerced_space_pairs"] = space_converted
    total = arrow_converted + space_converted
    if total:
        meta["coerced_pairs"] = total
    return sanitized, meta, has_section, leading_pairs


def _parse_ini_structured(
    text: str,
    *,
    has_section: bool,
    leading_pairs: bool,
) -> tuple[object, dict[str, object]]:
    parser = configparser.RawConfigParser(
        strict=False,
        interpolation=None,
        allow_no_value=True,
    )
    parser.optionxform = str
    synthetic_root = None
    try:
        if not has_section or leading_pairs:
            synthetic_root = "__root__"
            parser.read_string(f"[{synthetic_root}]\n{text}")
        else:
            parser.read_string(text)
    except configparser.Error as exc:
        raise ValueError("invalid ini data") from exc

    meta: dict[str, object] = {}

    if not has_section:
        assert synthetic_root is not None
        section = dict(parser._sections.get(synthetic_root, {}))  # type: ignore[attr-defined]
        section.pop("__name__", None)
        meta.update(
            {
                "sections": [],
                "section_count": 0,
                "key_count": len(section),
            }
        )
        return section, meta

    defaults: Dict[str, object]
    if synthetic_root is not None:
        defaults = dict(parser._sections.get(synthetic_root, {}))  # type: ignore[attr-defined]
        defaults.pop("__name__", None)
    else:
        defaults = dict(parser._defaults)  # type: ignore[attr-defined]

    sections: dict[str, dict[str, object]] = {}
    for name, payload in parser._sections.items():  # type: ignore[attr-defined]
        if synthetic_root is not None and name == synthetic_root:
            continue
        data = dict(payload)
        data.pop("__name__", None)
        sections[name] = data

    section_names = [name for name in parser.sections() if name != synthetic_root]

    ordered: Dict[str, object] = {}
    if defaults:
        ordered["<defaults>"] = defaults
    for name in section_names:
        ordered[name] = sections.get(name, {})

    meta.update(
        {
            "sections": section_names,
            "section_count": len(section_names),
        }
    )
    if defaults:
        meta["default_keys"] = list(defaults.keys())
    meta["key_count"] = sum(
        len(value) if isinstance(value, Mapping) else 1 for value in ordered.values()
    )
    return ordered, meta


def _structured_to_blocks(
    data: object,
    source: Path,
    *,
    path: PathTuple = (),
    limit: int = _STRUCTURED_BLOCK_LIMIT,
) -> Tuple[List[Block], bool]:
    blocks: List[Block] = []
    truncated = False

    def _add(block: Block) -> bool:
        nonlocal truncated
        if len(blocks) >= limit:
            truncated = True
            return False
        blocks.append(block)
        return True

    def _visit(value: object, cursor: PathTuple) -> None:
        if truncated:
            return

        label = _format_label(cursor)
        r_path = _format_r_path(cursor)
        glue_path = _format_glue_path(cursor)
        r_tokens = _format_r_tokens(cursor)

        if isinstance(value, Mapping):
            items = list(value.items())
            attrs = {
                "key": label,
                "type": "object",
                "size": len(items),
                "keys": [str(key) for key, _ in items[:20]],
                "path_r": r_path,
                "path_r_tokens": r_tokens,
                "path_glue": glue_path,
            }
            sample_values = [
                _stringify_scalar(child)
                for _, child in items
                if not isinstance(child, Mapping)
                and not (
                    isinstance(child, Sequence)
                    and not isinstance(child, (str, bytes, bytearray))
                )
            ][:5]
            if sample_values:
                attrs["sample_values"] = sample_values
            if not _add(
                Block(
                    type="metadata",
                    text=f"{label}: object ({len(items)} keys)",
                    attrs=attrs,
                    source=str(source),
                    confidence=0.7,
                )
            ):
                return
            for key, child in items:
                child_path = _extend_key(cursor, key)
                _visit(child, child_path)
            return

        if isinstance(value, Sequence) and not isinstance(value, (str, bytes, bytearray)):
            seq = list(value)
            sample_types = sorted({type(item).__name__ for item in seq[:5]})
            attrs = {
                "key": label,
                "type": "array",
                "size": len(seq),
                "sample_types": sample_types,
                "path_r": r_path,
                "path_r_tokens": r_tokens,
                "path_glue": glue_path,
            }
            sample_values = [
                _stringify_scalar(item)
                for item in seq
                if not isinstance(item, Mapping)
                and not (
                    isinstance(item, Sequence)
                    and not isinstance(item, (str, bytes, bytearray))
                )
            ][:5]
            if sample_values:
                attrs["sample_values"] = sample_values
            if not _add(
                Block(
                    type="metadata",
                    text=f"{label}: array ({len(seq)} items)",
                    attrs=attrs,
                    source=str(source),
                    confidence=0.65,
                )
            ):
                return
            for idx, child in enumerate(seq):
                child_path = _extend_index(cursor, idx, base=0)
                _visit(child, child_path)
            return

        value_text = _stringify_scalar(value)
        text = value_text if not cursor else f"{label}: {value_text}"
        attrs = {
            "key": label,
            "value": value_text,
            "value_type": type(value).__name__,
            "path_r": r_path,
            "path_r_tokens": r_tokens,
            "path_glue": glue_path,
        }
        _add(
            Block(
                type="kv",
                text=text,
                attrs=attrs,
                source=str(source),
                confidence=0.85,
            )
        )

    _visit(data, path)
    return blocks, truncated


def parse_txt(path: str | Path) -> List[Block]:
    source = Path(path)
    text = source.read_text(encoding="utf-8", errors="ignore")
    blocks: List[Block] = []
    for chunk in _iter_refined_chunks(text):
        if "\n" in chunk:
            lines = [line for line in chunk.splitlines() if line.strip()]
            result = _coerce_tabular_rows(lines)
            if result is not None:
                delimiter, rows = result
                blocks.append(
                    Block(
                        type="table",
                        text="\n".join([", ".join(row) for row in rows]),
                        attrs={
                            "rows": json.dumps(rows, ensure_ascii=False),
                            "delimiter": delimiter,
                            "structured_from": "text_table",
                            "row_count": len(rows),
                            "column_count": len(rows[0]) if rows else 0,
                        },
                        source=str(source),
                        confidence=0.78 if delimiter != "whitespace" else 0.74,
                    )
                )
                continue
        blocks.append(_block_from_chunk(chunk, source))
    return blocks or [_new_block("other", text, source, confidence=0.3)]


def parse_md(path: str | Path) -> List[Block]:
    source = Path(path)
    text = source.read_text(encoding="utf-8", errors="ignore")
    blocks: List[Block] = []
    in_code = False
    code_buffer: List[str] = []
    for line in text.splitlines():
        stripped = line.strip("\n")
        if stripped.startswith("```"):
            if in_code:
                blocks.append(
                    _new_block("code", "\n".join(code_buffer), source, confidence=0.7)
                )
                code_buffer.clear()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue
        if stripped.startswith("#"):
            blocks.append(_new_block("heading", stripped.lstrip("# "), source, confidence=0.8))
            continue
        if _LIST_MARKERS.match(stripped):
            blocks.append(_block_from_chunk(stripped, source))
            continue
        if stripped:
            for chunk in _shatter_chunk(stripped):
                blocks.append(_block_from_chunk(chunk, source))
    if code_buffer:
        blocks.append(_new_block("code", "\n".join(code_buffer), source, confidence=0.7))
    return blocks or [_new_block("other", text, source, confidence=0.3)]


def parse_html(path: str | Path) -> List[Block]:
    source = Path(path)
    html = source.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    blocks: List[Block] = []
    for element in soup.find_all(["title", "h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "code", "table"]):
        text = element.get_text(" ", strip=True)
        if not text:
            continue
        name = element.name or "p"
        if name == "title":
            blocks.append(_new_block("title", text, source, confidence=0.9))
        elif name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            blocks.append(_new_block("heading", text, source, confidence=0.85))
        elif name == "li":
            blocks.append(_new_block("list_item", text, source, confidence=0.7))
        elif name in {"pre", "code"}:
            blocks.append(_new_block("code", text, source, confidence=0.75))
        elif name == "table":
            rows = [
                [cell.get_text(" ", strip=True) for cell in row.find_all(["th", "td"])]
                for row in element.find_all("tr")
            ]
            blocks.append(
                Block(
                    type="table",
                    text="\n".join([", ".join(row) for row in rows]),
                    attrs={"rows": json.dumps(rows, ensure_ascii=False)},
                    source=str(source),
                    confidence=0.8,
                )
            )
        else:
            for chunk in _shatter_chunk(text):
                blocks.append(_block_from_chunk(chunk, source))
    return blocks or [_new_block("paragraph", soup.get_text("\n", strip=True), source, confidence=0.4)]


def parse_csv(path: str | Path) -> List[Block]:
    source = Path(path)
    with source.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        reader = csv.reader(handle)
        rows = [row for row in reader]
    text = "\n".join([", ".join(row) for row in rows])
    return [
        Block(
            type="table",
            text=text,
            attrs={"rows": json.dumps(rows, ensure_ascii=False)},
            source=str(source),
            confidence=0.8,
        )
    ]


def stream_pdf(path: str | Path) -> Iterator[Block]:
    source = Path(path)
    reader = PdfReader(str(source))
    analyzer = VisualLayoutAnalyzer(profile="pdf")
    emitted = False
    for index, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        candidates: List[LayoutCandidate] = []
        for order, chunk in enumerate(_iter_refined_chunks(text)):
            base = _block_from_chunk(chunk, source)
            attrs = dict(base.attrs)
            attrs.setdefault("page", str(index + 1))
            base_block = clone_model(base, attrs=attrs, confidence=max(base.confidence, 0.55))
            span_height = min(160.0, 24.0 + len(chunk) * 0.12)
            span_width = min(640.0, 120.0 + len(chunk) * 0.9)
            top = float(order) * (span_height + 6.0)
            bbox = (36.0, top, 36.0 + span_width, top + span_height)
            score = max(base_block.confidence, min(0.95, 0.35 + len(chunk) / 600.0))
            candidates.append(
                LayoutCandidate(
                    block=base_block,
                    bbox=bbox,
                    page=index,
                    score=score,
                    order_hint=order,
                    metadata={"layout_source": "pdf"},
                )
            )
        for segment in analyzer.process(candidates):
            emitted = True
            yield segment.block
    if not emitted:
        yield _new_block("other", "", source, confidence=0.1)


def parse_pdf(path: str | Path) -> List[Block]:
    return list(stream_pdf(path))


def stream_image(path: str | Path) -> Iterator[Block]:
    source = Path(path)
    text, meta, segments = _extract_image_text(source)

    analyzer = VisualLayoutAnalyzer(profile="image")
    emitted = 0
    truncated = False

    summary_attrs = dict(meta)
    if summary_attrs:
        yield Block(
            type="metadata",
            text="Image summary",
            attrs=summary_attrs,
            source=str(source),
            confidence=0.45,
        )
        emitted += 1

    candidates: List[LayoutCandidate] = []
    metadata_pending: List[Block] = []
    for index, segment in enumerate(segments):
        if emitted + len(candidates) >= _STRUCTURED_BLOCK_LIMIT:
            truncated = True
            break
        seg_text = str(segment.get("text", "")).strip()
        kind = segment.get("kind", "ocr")
        source_kind = segment.get("source", "ocr")
        if kind == "metadata":
            attrs = {"image_source": source_kind, "length": len(seg_text)}
            key = segment.get("key")
            if key:
                attrs["key"] = key
            metadata_pending.append(
                Block(
                    type="metadata",
                    text=seg_text,
                    attrs=attrs,
                    source=str(source),
                    confidence=0.6,
                )
            )
            continue
        if not seg_text:
            continue
        prov = Provenance()
        order = segment.get("order")
        if order is not None:
            try:
                prov.order = int(order)
            except Exception:
                prov.order = None
        page = segment.get("page")
        if page is not None:
            try:
                prov.page = int(page)
            except Exception:
                prov.page = None
        bbox_raw = segment.get("bbox")
        bbox: Sequence[float] | None = None
        if bbox_raw and isinstance(bbox_raw, Sequence) and len(bbox_raw) == 4:
            try:
                prov.bbox = BBox(
                    x0=float(bbox_raw[0]),
                    y0=float(bbox_raw[1]),
                    x1=float(bbox_raw[2]),
                    y1=float(bbox_raw[3]),
                )
                bbox = tuple(float(v) for v in bbox_raw)
            except Exception:
                prov.bbox = None
        attrs: Dict[str, Any] = {"image_source": source_kind}
        if source_kind == "ocr":
            languages = meta.get("image_ocr_languages")
            if languages:
                attrs["ocr_languages"] = languages
        confidence = segment.get("confidence", 0.6)
        try:
            conf_value = float(confidence)
            attrs["confidence"] = conf_value
        except Exception:
            attrs["confidence"] = confidence
            conf_value = 0.6
        base_block = Block(
            type="paragraph",
            text=seg_text,
            attrs=attrs,
            prov=prov,
            source=str(source),
            confidence=conf_value,
        )
        order_hint = prov.order if prov.order is not None else index
        if bbox is None:
            top = float(order_hint) * 30.0
            bbox = (12.0, top, 12.0 + min(600.0, 80.0 + len(seg_text) * 1.2), top + 24.0)
        candidates.append(
            LayoutCandidate(
                block=base_block,
                bbox=bbox,
                page=prov.page or 0,
                score=conf_value,
                order_hint=order_hint,
                metadata={"image_source": source_kind, "layout_source": "image"},
            )
        )

    for meta_block in metadata_pending:
        if emitted >= _STRUCTURED_BLOCK_LIMIT:
            truncated = True
            break
        yield meta_block
        emitted += 1

    for segment in analyzer.process(candidates):
        if emitted >= _STRUCTURED_BLOCK_LIMIT:
            truncated = True
            break
        yield segment.block
        emitted += 1

    if emitted == 0 and text:
        for chunk in _iter_refined_chunks(text):
            if emitted >= _STRUCTURED_BLOCK_LIMIT:
                truncated = True
                break
            yield _block_from_chunk(chunk, source)
            emitted += 1

    if emitted == 0:
        yield Block(
            type="metadata",
            text="Image contained no extractable text",
            attrs={"image_has_text": bool(text)},
            source=str(source),
            confidence=0.35,
        )
        emitted += 1

    if truncated and emitted < _STRUCTURED_BLOCK_LIMIT:
        yield Block(
            type="metadata",
            text=f"Truncated after {_STRUCTURED_BLOCK_LIMIT - 1} structured blocks",
            attrs={"truncated": True},
            source=str(source),
            confidence=0.4,
        )


def parse_image(path: str | Path) -> List[Block]:
    return list(stream_image(path))[:_STRUCTURED_BLOCK_LIMIT]


def parse_docx(path: str | Path) -> List[Block]:
    source = Path(path)
    doc = DocxDocument(str(source))
    blocks: List[Block] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name.lower() if paragraph.style and paragraph.style.name else "")
        is_heading = "heading" in style if style else False
        for chunk in _shatter_chunk(text):
            base = _block_from_chunk(chunk, source)
            data: Dict[str, object] = {"confidence": max(base.confidence, 0.65)}
            if is_heading and base.type == "paragraph":
                data["type"] = "heading"
            blocks.append(clone_model(base, **data))
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        blocks.append(
            Block(
                type="table",
                text="\n".join([", ".join(row) for row in rows]),
                attrs={"rows": json.dumps(rows, ensure_ascii=False)},
                source=str(source),
                confidence=0.75,
            )
        )
    return blocks or [_new_block("other", "", source, confidence=0.1)]


def parse_pptx(path: str | Path) -> List[Block]:
    source = Path(path)
    blocks: List[Block] = []
    try:
        with zipfile.ZipFile(str(source)) as zf:
            slide_names = sorted(
                name
                for name in zf.namelist()
                if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            )
            for index, slide_name in enumerate(slide_names, 1):
                try:
                    root = ET.fromstring(zf.read(slide_name))
                except Exception:
                    continue
                texts: List[str] = []
                for node in root.iter():
                    if node.tag.endswith("}t") and (node.text or "").strip():
                        texts.append(node.text.strip())
                if not texts:
                    continue
                slide_text = "\n".join(texts)
                for chunk in _iter_refined_chunks(slide_text):
                    base = _block_from_chunk(chunk, source)
                    attrs = dict(base.attrs)
                    attrs.setdefault("slide", str(index))
                    blocks.append(
                        clone_model(base, attrs=attrs, confidence=max(base.confidence, 0.6))
                    )
    except Exception:
        pass
    return blocks or [_new_block("other", "", source, confidence=0.2)]


def parse_xlsx(path: str | Path) -> List[Block]:
    source = Path(path)
    workbook = load_workbook(filename=str(source), read_only=True, data_only=True)
    sheet = workbook.active
    rows = [
        ["" if cell.value is None else str(cell.value) for cell in row]
        for row in sheet.iter_rows(values_only=True)
    ]
    text = "\n".join([", ".join(row) for row in rows])
    workbook.close()
    return [
        Block(
            type="table",
            text=text,
            attrs={"rows": json.dumps(rows, ensure_ascii=False)},
            source=str(source),
            confidence=0.75,
        )
    ]


def parse_json(path: str | Path) -> List[Block]:
    source = Path(path)
    raw = source.read_text(encoding="utf-8", errors="ignore")
    coercion_origin: str | None = None
    try:
        data, coercion = _load_json_like(raw)
    except ValueError:
        return parse_txt(path)
    else:
        coercion_origin = coercion

    budget = max(1, _STRUCTURED_BLOCK_LIMIT - 1)
    blocks, truncated = _structured_to_blocks(data, source, limit=budget)
    if not blocks:
        return parse_txt(path)
    if coercion_origin:
        for block in blocks:
            key = block.attrs.get("key") if block.attrs else None
            if key == "<root>" or block is blocks[0]:
                block.attrs = dict(block.attrs)
                block.attrs["coerced_from"] = coercion_origin
                break
    if truncated and len(blocks) < _STRUCTURED_BLOCK_LIMIT:
        blocks.append(
            Block(
                type="metadata",
                text=f"Truncated after {_STRUCTURED_BLOCK_LIMIT - 1} structured blocks",
                attrs={
                    "truncated": True,
                    "key": _format_label(()),
                    "path_r": _format_r_path(()),
                    "path_r_tokens": _format_r_tokens(()),
                    "path_glue": _format_glue_path(()),
                },
                source=str(source),
                confidence=0.4,
            )
        )
    return blocks


def parse_ini(path: str | Path) -> List[Block]:
    source = Path(path)
    raw = source.read_text(encoding="utf-8", errors="ignore")
    sanitized, coercion_meta, has_section, leading_pairs = _prepare_ini_input(raw)
    try:
        data, structure_meta = _parse_ini_structured(
            sanitized, has_section=has_section, leading_pairs=leading_pairs
        )
    except ValueError:
        return parse_txt(path)

    info: Dict[str, object] = {}
    info.update(coercion_meta)
    info.update(structure_meta)

    budget = max(1, _STRUCTURED_BLOCK_LIMIT - 1)
    blocks, truncated = _structured_to_blocks(data, source, limit=budget)
    if not blocks:
        return parse_txt(path)

    if info:
        head = blocks[0]
        head.attrs = dict(head.attrs)
        head.attrs.update(info)

    if truncated and len(blocks) < _STRUCTURED_BLOCK_LIMIT:
        blocks.append(
            Block(
                type="metadata",
                text=f"Truncated after {_STRUCTURED_BLOCK_LIMIT - 1} structured blocks",
                attrs={
                    "truncated": True,
                    "key": _format_label(()),
                    "path_r": _format_r_path(()),
                    "path_r_tokens": _format_r_tokens(()),
                    "path_glue": _format_glue_path(()),
                },
                source=str(source),
                confidence=0.4,
            )
        )

    return blocks


def parse_jsonl(path: str | Path) -> List[Block]:
    source = Path(path)
    text = source.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()
    budget = max(1, _STRUCTURED_BLOCK_LIMIT - 1)
    blocks: List[Block] = []
    truncated = False

    def _try_add(block: Block) -> bool:
        nonlocal truncated
        if len(blocks) >= budget:
            truncated = True
            return False
        blocks.append(block)
        return True

    for idx, line in enumerate(lines, 1):
        stripped = line.strip()
        if not stripped:
            continue
        if len(blocks) >= budget:
            truncated = True
            break
        try:
            record = json.loads(stripped)
        except json.JSONDecodeError:
            if not _try_add(_new_block("other", stripped, source, confidence=0.25)):
                break
            continue

        record_path = _extend_index(_extend_key((), "record"), idx, base=1)
        label = _format_label(record_path)
        summary_type = type(record).__name__
        summary_attrs = {
            "key": label,
            "type": summary_type,
            "line": idx,
            "path_r": _format_r_path(record_path),
            "path_r_tokens": _format_r_tokens(record_path),
            "path_glue": _format_glue_path(record_path),
        }
        if isinstance(record, Mapping):
            summary_attrs["size"] = len(record)
            summary_attrs["keys"] = [str(k) for k in list(record.keys())[:10]]
        elif isinstance(record, Sequence) and not isinstance(record, (str, bytes, bytearray)):
            summary_attrs["size"] = len(record)
        if not _try_add(
            Block(
                type="metadata",
                text=f"{label}: {summary_type}",
                attrs=summary_attrs,
                source=str(source),
                confidence=0.65,
            )
        ):
            break

        remaining = budget - len(blocks)
        if remaining <= 0:
            truncated = True
            break
        sub_blocks, sub_truncated = _structured_to_blocks(record, source, path=record_path, limit=remaining)
        blocks.extend(sub_blocks)
        if sub_truncated:
            truncated = True
            break

    if not blocks:
        return parse_txt(path)
    if truncated and len(blocks) < _STRUCTURED_BLOCK_LIMIT:
        blocks.append(
            Block(
                type="metadata",
                text=f"Truncated after {_STRUCTURED_BLOCK_LIMIT - 1} structured blocks",
                attrs={
                    "truncated": True,
                    "key": _format_label(()),
                    "path_r": _format_r_path(()),
                    "path_r_tokens": _format_r_tokens(()),
                    "path_glue": _format_glue_path(()),
                },
                source=str(source),
                confidence=0.4,
            )
        )
    return blocks


def parse_yaml(path: str | Path) -> List[Block]:
    source = Path(path)
    raw = source.read_text(encoding="utf-8", errors="ignore")
    try:
        docs = list(yaml.safe_load_all(raw))
    except Exception:
        return parse_txt(path)

    if not docs:
        return parse_txt(path)

    budget = max(1, _STRUCTURED_BLOCK_LIMIT - 1)
    blocks: List[Block] = []
    truncated = False

    def _try_add(block: Block) -> bool:
        nonlocal truncated
        if len(blocks) >= budget:
            truncated = True
            return False
        blocks.append(block)
        return True

    sample_types = [type(doc).__name__ for doc in docs[:5]]
    _try_add(
        Block(
            type="metadata",
            text=f"YAML: {len(docs)} document(s)",
            attrs={
                "key": _format_label(()),
                "documents": len(docs),
                "sample_types": sample_types,
                "path_r": _format_r_path(()),
                "path_r_tokens": _format_r_tokens(()),
                "path_glue": _format_glue_path(()),
            },
            source=str(source),
            confidence=0.6,
        )
    )

    for idx, doc in enumerate(docs):
        if len(blocks) >= budget:
            truncated = True
            break
        remaining = budget - len(blocks)
        if remaining <= 0:
            truncated = True
            break
        doc_path: PathTuple
        if len(docs) == 1:
            doc_path = ()
        else:
            doc_path = _extend_index(_extend_key((), "doc"), idx, base=0)
        sub_blocks, sub_truncated = _structured_to_blocks(doc, source, path=doc_path, limit=remaining)
        blocks.extend(sub_blocks)
        if sub_truncated:
            truncated = True
            break

    if not blocks:
        return parse_txt(path)
    if truncated and len(blocks) < _STRUCTURED_BLOCK_LIMIT:
        blocks.append(
            Block(
                type="metadata",
                text=f"Truncated after {_STRUCTURED_BLOCK_LIMIT - 1} structured blocks",
                attrs={
                    "truncated": True,
                    "key": _format_label(()),
                    "path_r": _format_r_path(()),
                    "path_r_tokens": _format_r_tokens(()),
                    "path_glue": _format_glue_path(()),
                },
                source=str(source),
                confidence=0.4,
            )
        )
    return blocks


def parse_toml(path: str | Path) -> List[Block]:
    source = Path(path)
    raw = source.read_text(encoding="utf-8", errors="ignore")
    try:
        try:
            import tomllib  # type: ignore[attr-defined]
        except ModuleNotFoundError:  # pragma: no cover - Python <3.11 fallback
            import tomli as tomllib  # type: ignore
    except ModuleNotFoundError:
        return parse_txt(path)

    try:
        data = tomllib.loads(raw)
    except Exception:
        return parse_txt(path)

    budget = max(1, _STRUCTURED_BLOCK_LIMIT - 1)
    blocks, truncated = _structured_to_blocks(data, source, limit=budget)
    if not blocks:
        return parse_txt(path)
    if truncated and len(blocks) < _STRUCTURED_BLOCK_LIMIT:
        blocks.append(
            Block(
                type="metadata",
                text=f"Truncated after {_STRUCTURED_BLOCK_LIMIT - 1} structured blocks",
                attrs={
                    "truncated": True,
                    "key": _format_label(()),
                    "path_r": _format_r_path(()),
                    "path_r_tokens": _format_r_tokens(()),
                    "path_glue": _format_glue_path(()),
                },
                source=str(source),
                confidence=0.4,
            )
        )
    return blocks


def parse_eml(path: str | Path) -> List[Block]:
    source = Path(path)
    data = source.read_bytes()
    blocks: List[Block] = []
    try:
        message = BytesParser(policy=policy.default).parsebytes(data)
    except Exception:
        return [_new_block("other", data.decode("utf-8", errors="ignore"), source, confidence=0.2)]

    header_attrs = {
        "subject": message.get("subject", ""),
        "from": message.get("from", ""),
        "to": message.get("to", ""),
        "cc": message.get("cc", ""),
        "date": message.get("date", ""),
    }
    header_text = "\n".join(
        [f"Subject: {header_attrs['subject']}", f"From: {header_attrs['from']}", f"To: {header_attrs['to']}"]
    ).strip()
    blocks.append(
        Block(
            type="meta",
            text=header_text,
            attrs={k: v for k, v in header_attrs.items() if v},
            source=str(source),
            confidence=0.85,
        )
    )

    text_parts: List[str] = []
    html_fallback: List[str] = []
    attachments: List[Dict[str, object]] = []
    for part in message.walk():
        if part.is_multipart():
            continue
        filename = part.get_filename()
        content_type = part.get_content_type()
        payload = part.get_payload(decode=True) or b""
        if filename:
            attachments.append(
                {
                    "filename": filename,
                    "content_type": content_type,
                    "size": len(payload),
                }
            )
            continue
        if content_type.startswith("text/"):
            charset = part.get_content_charset() or "utf-8"
            try:
                decoded = payload.decode(charset, errors="ignore")
            except LookupError:
                decoded = payload.decode("utf-8", errors="ignore")
            if content_type == "text/plain":
                text_parts.append(decoded)
            else:
                html_fallback.append(decoded)

    if not text_parts and html_fallback:
        for html in html_fallback:
            cleaned = BeautifulSoup(html, "html.parser").get_text(" ", strip=True)
            if cleaned:
                text_parts.append(cleaned)

    for part in text_parts:
        for chunk in _iter_refined_chunks(part):
            blocks.append(_block_from_chunk(chunk, source))

    for attachment in attachments[:10]:
        blocks.append(
            Block(
                type="attachment",
                text=attachment.get("filename") or attachment.get("content_type", "attachment"),
                attrs=attachment,
                source=str(source),
                confidence=0.4,
            )
        )

    return blocks or [_new_block("other", "", source, confidence=0.2)]


def parse_ics(path: str | Path) -> List[Block]:
    source = Path(path)
    raw_lines = source.read_text(encoding="utf-8", errors="ignore").splitlines()
    unfolded: List[str] = []
    for line in raw_lines:
        if line.startswith((" ", "\t")) and unfolded:
            unfolded[-1] += line.strip()
        else:
            unfolded.append(line)

    events: List[Dict[str, str]] = []
    current: Dict[str, str] = {}
    for line in unfolded:
        upper = line.upper()
        if upper == "BEGIN:VEVENT":
            current = {}
            continue
        if upper == "END:VEVENT":
            if current:
                events.append(current)
            current = {}
            continue
        if ":" in line:
            key, value = line.split(":", 1)
            current[key.strip().upper()] = value.strip()

    blocks: List[Block] = []
    for event in events:
        summary = event.get("SUMMARY", "Event")
        description_lines = [summary]
        if "DTSTART" in event:
            description_lines.append(f"Start: {event['DTSTART']}")
        if "DTEND" in event:
            description_lines.append(f"End: {event['DTEND']}")
        if "LOCATION" in event:
            description_lines.append(f"Location: {event['LOCATION']}")
        if "DESCRIPTION" in event:
            description_lines.append(event["DESCRIPTION"])
        blocks.append(
            Block(
                type="event",
                text="\n".join(description_lines),
                attrs={k.lower(): v for k, v in event.items()},
                source=str(source),
                confidence=0.75,
            )
        )

    if not blocks:
        text = "\n".join(unfolded)
        blocks.append(_new_block("other", text, source, confidence=0.3))

    return blocks

